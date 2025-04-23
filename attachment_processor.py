import os
import logging
import tempfile
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import mimetypes
import re

# Importações para processamento de diferentes tipos de arquivos
import PyPDF2
import pandas as pd
from docx import Document
import csv
import json
from bs4 import BeautifulSoup

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("attachment_processor")

class AttachmentProcessorInterface:
    """Interface para processadores de anexos."""
    
    def can_process(self, content_type: str, filename: str) -> bool:
        """
        Verifica se o processador pode processar o tipo de anexo.
        
        Args:
            content_type: Tipo MIME do anexo
            filename: Nome do arquivo
            
        Returns:
            bool: True se o processador pode processar o anexo, False caso contrário
        """
        pass
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Processa um anexo e extrai seu conteúdo.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dict[str, Any]: Dicionário com conteúdo extraído e metadados
        """
        pass

class PDFProcessor(AttachmentProcessorInterface):
    """Processador para arquivos PDF."""
    
    def can_process(self, content_type: str, filename: str) -> bool:
        """
        Verifica se o anexo é um PDF.
        
        Args:
            content_type: Tipo MIME do anexo
            filename: Nome do arquivo
            
        Returns:
            bool: True se o anexo é um PDF, False caso contrário
        """
        return content_type == 'application/pdf' or filename.lower().endswith('.pdf')
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai texto de um arquivo PDF.
        
        Args:
            file_path: Caminho para o arquivo PDF
            
        Returns:
            Dict[str, Any]: Dicionário com texto extraído e metadados
        """
        try:
            text = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extrai metadados
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', '')
                    }
                
                # Extrai texto de todas as páginas
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            
            # Procura por padrões específicos no texto (números de contrato, nomes de escolas, etc.)
            contract_numbers = self._extract_contract_numbers(text)
            school_names = self._extract_school_names(text)
            dates = self._extract_dates(text)
            
            return {
                'content_type': 'text/plain',
                'text': text,
                'metadata': metadata,
                'contract_numbers': contract_numbers,
                'school_names': school_names,
                'dates': dates,
                'page_count': len(pdf_reader.pages)
            }
        
        except Exception as e:
            logger.error(f"Erro ao processar PDF {file_path}: {str(e)}")
            return {
                'content_type': 'text/plain',
                'text': f"[Erro ao processar PDF: {str(e)}]",
                'metadata': {},
                'error': str(e)
            }
    
    def _extract_contract_numbers(self, text: str) -> List[str]:
        """
        Extrai possíveis números de contrato do texto.
        
        Args:
            text: Texto do PDF
            
        Returns:
            List[str]: Lista de possíveis números de contrato
        """
        # Padrões comuns para números de contrato
        patterns = [
            r'contrato\s+n[º°]?\s*[:.]?\s*(\d{4,10}[-/]?[\dA-Z]{0,5})',  # Contrato Nº: 12345
            r'contrato\s+(\d{4,10}[-/]?[\dA-Z]{0,5})',                   # Contrato 12345
            r'n[º°]?\s*[:.]?\s*(\d{4,10}[-/]?[\dA-Z]{0,5})',             # Nº: 12345
            r'processo\s+n[º°]?\s*[:.]?\s*(\d{4,10}[-/]?[\dA-Z]{0,5})'   # Processo Nº: 12345
        ]
        
        contract_numbers = []
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                contract_numbers.append(match.group(1))
        
        return list(set(contract_numbers))  # Remove duplicatas
    
    def _extract_school_names(self, text: str) -> List[str]:
        """
        Extrai possíveis nomes de escolas do texto.
        
        Args:
            text: Texto do PDF
            
        Returns:
            List[str]: Lista de possíveis nomes de escolas
        """
        # Padrões comuns para nomes de escolas
        patterns = [
            r'escola\s+([A-Za-zÀ-ÖØ-öø-ÿ\s]{3,50})',                     # Escola Nome da Escola
            r'colégio\s+([A-Za-zÀ-ÖØ-öø-ÿ\s]{3,50})',                    # Colégio Nome do Colégio
            r'centro\s+educacional\s+([A-Za-zÀ-ÖØ-öø-ÿ\s]{3,50})',       # Centro Educacional Nome
            r'e\.e\.\s+([A-Za-zÀ-ÖØ-öø-ÿ\s]{3,50})',                     # E.E. Nome da Escola
            r'e\.m\.\s+([A-Za-zÀ-ÖØ-öø-ÿ\s]{3,50})'                      # E.M. Nome da Escola
        ]
        
        school_names = []
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Limpa o nome da escola (remove espaços extras, pontuação no final, etc.)
                school_name = match.group(1).strip()
                school_name = re.sub(r'[.,;:]$', '', school_name)
                if len(school_name) > 3:  # Ignora nomes muito curtos
                    school_names.append(school_name)
        
        return list(set(school_names))  # Remove duplicatas
    
    def _extract_dates(self, text: str) -> List[str]:
        """
        Extrai possíveis datas do texto.
        
        Args:
            text: Texto do PDF
            
        Returns:
            List[str]: Lista de possíveis datas
        """
        # Padrões comuns para datas
        patterns = [
            r'(\d{2}/\d{2}/\d{4})',                                      # DD/MM/AAAA
            r'(\d{2}\.\d{2}\.\d{4})',                                     # DD.MM.AAAA
            r'(\d{2}-\d{2}-\d{4})',                                       # DD-MM-AAAA
            r'(\d{1,2}\s+de\s+[a-zç]+\s+de\s+\d{4})'                     # DD de Mês de AAAA
        ]
        
        dates = []
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                dates.append(match.group(1))
        
        return list(set(dates))  # Remove duplicatas

class DocxProcessor(AttachmentProcessorInterface):
    """Processador para arquivos DOCX."""
    
    def can_process(self, content_type: str, filename: str) -> bool:
        """
        Verifica se o anexo é um DOCX.
        
        Args:
            content_type: Tipo MIME do anexo
            filename: Nome do arquivo
            
        Returns:
            bool: True se o anexo é um DOCX, False caso contrário
        """
        return content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or filename.lower().endswith('.docx')
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai texto de um arquivo DOCX.
        
        Args:
            file_path: Caminho para o arquivo DOCX
            
        Returns:
            Dict[str, Any]: Dicionário com texto extraído e metadados
        """
        try:
            doc = Document(file_path)
            
            # Extrai texto de parágrafos
            paragraphs = [p.text for p in doc.paragraphs]
            text = "\n".join(paragraphs)
            
            # Extrai texto de tabelas
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))
            
            if tables_text:
                text += "\n\n--- TABELAS ---\n\n" + "\n".join(tables_text)
            
            # Extrai propriedades do documento
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
            }
            
            # Procura por padrões específicos no texto
            contract_numbers = self._extract_contract_numbers(text)
            school_names = self._extract_school_names(text)
            dates = self._extract_dates(text)
            
            return {
                'content_type': 'text/plain',
                'text': text,
                'metadata': metadata,
                'contract_numbers': contract_numbers,
                'school_names': school_names,
                'dates': dates
            }
        
        except Exception as e:
            logger.error(f"Erro ao processar DOCX {file_path}: {str(e)}")
            return {
                'content_type': 'text/plain',
                'text': f"[Erro ao processar DOCX: {str(e)}]",
                'metadata': {},
                'error': str(e)
            }
    
    # Reutiliza os métodos de extração do PDFProcessor
    def _extract_contract_numbers(self, text: str) -> List[str]:
        return PDFProcessor()._extract_contract_numbers(text)
    
    def _extract_school_names(self, text: str) -> List[str]:
        return PDFProcessor()._extract_school_names(text)
    
    def _extract_dates(self, text: str) -> List[str]:
        return PDFProcessor()._extract_dates(text)

class ExcelProcessor(AttachmentProcessorInterface):
    """Processador para arquivos Excel (XLS, XLSX)."""
    
    def can_process(self, content_type: str, filename: str) -> bool:
        """
        Verifica se o anexo é um arquivo Excel.
        
        Args:
            content_type: Tipo MIME do anexo
            filename: Nome do arquivo
            
        Returns:
            bool: True se o anexo é um arquivo Excel, False caso contrário
        """
        excel_types = [
            'application/vnd.ms-excel',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel.sheet.macroEnabled.12'
        ]
        return content_type in excel_types or filename.lower().endswith(('.xls', '.xlsx', '.xlsm'))
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai dados de um arquivo Excel.
        
        Args:
            file_path: Caminho para o arquivo Excel
            
        Returns:
            Dict[str, Any]: Dicionário com dados extraídos e metadados
        """
        try:
            # Lê todas as planilhas
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            # Converte cada planilha para texto
            sheets_text = []
            sheets_data = {}
            
            for sheet_name, df in excel_data.items():
                # Converte para texto
                sheet_text = f"--- PLANILHA: {sheet_name} ---\n"
                sheet_text += df.to_string(index=False) + "\n\n"
                sheets_text.append(sheet_text)
                
                # Armazena dados estruturados
                sheets_data[sheet_name] = df.to_dict(orient='records')
            
            text = "\n".join(sheets_text)
            
            # Procura por padrões específicos no texto
            contract_numbers = self._extract_contract_numbers(text)
            school_names = self._extract_school_names(text)
            dates = self._extract_dates(text)
            
            return {
                'content_type': 'application/json',
                'text': text,
                'structured_data': sheets_data,
                'metadata': {
                    'sheet_count': len(excel_data),
                    'sheet_names': list(excel_data.keys())
                },
                'contract_numbers': contract_numbers,
                'school_names': school_names,
                'dates': dates
            }
        
        except Exception as e:
            logger.error(f"Erro ao processar Excel {file_path}: {str(e)}")
            return {
                'content_type': 'text/plain',
                'text': f"[Erro ao processar Excel: {str(e)}]",
                'metadata': {},
                'error': str(e)
            }
    
    # Reutiliza os métodos de extração do PDFProcessor
    def _extract_contract_numbers(self, text: str) -> List[str]:
        return PDFProcessor()._extract_contract_numbers(text)
    
    def _extract_school_names(self, text: str) -> List[str]:
        return PDFProcessor()._extract_school_names(text)
    
    def _extract_dates(self, text: str) -> List[str]:
        return PDFProcessor()._extract_dates(text)

class CSVProcessor(AttachmentProcessorInterface):
    """Processador para arquivos CSV."""
    
    def can_process(self, content_type: str, filename: str) -> bool:
        """
        Verifica se o anexo é um arquivo CSV.
        
        Args:
            content_type: Tipo MIME do anexo
            filename: Nome do arquivo
            
        Returns:
            bool: True se o anexo é um arquivo CSV, False caso contrário
        """
        return content_type == 'text/csv' or filename.lower().endswith('.csv')
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai dados de um arquivo CSV.
        
        Args:
            file_path: Caminho para o arquivo CSV
            
        Returns:
            Dict[str, Any]: Dicionário com dados extraídos e metadados
        """
        try:
            # Tenta diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Não foi possível decodificar o arquivo CSV com os encodings disponíveis")
            
            # Converte para texto
            text = df.to_string(index=False)
            
            # Converte para dados estruturados
            data = df.to_dict(orient='records')
            
            # Procura por padrões específicos no texto
            contract_numbers = self._extract_contract_numbers(text)
            school_names = self._extract_school_names(text)
            dates = self._extract_dates(text)
            
            return {
                'content_type': 'application/json',
                'text': text,
                'structured_data': data,
                'metadata': {
                    'column_count': len(df.columns),
                    'row_count': len(df),
                    'columns': list(df.columns)
                },
                'contract_numbers': contract_numbers,
                'school_names': school_names,
                'dates': dates
            }
        
        except Exception as e:
            logger.error(f"Erro ao processar CSV {file_path}: {str(e)}")
            return {
                'content_type': 'text/plain',
                'text': f"[Erro ao processar CSV: {str(e)}]",
                'metadata': {},
                'error': str(e)
            }
    
    # Reutiliza os métodos de extração do PDFProcessor
    def _extract_contract_numbers(self, text: str) -> List[str]:
        return PDFProcessor()._extract_contract_numbers(text)
    
    def _extract_school_names(self, text: str) -> List[str]:
        return PDFProcessor()._extract_school_names(text)
    
    def _extract_dates(self, text: str) -> List[str]:
        return PDFProcessor()._extract_dates(text)

class TextProcessor(AttachmentProcessorInterface):
    """Processador para arquivos de texto."""
    
    def can_process(self, content_type: str, filename: str) -> bool:
        """
        Verifica se o anexo é um arquivo de texto.
        
        Args:
            content_type: Tipo MIME do anexo
            filename: Nome do arquivo
            
        Returns:
            bool: True se o anexo é um arquivo de texto, False caso contrário
        """
        text_types = [
            'text/plain',
            'text/html',
            'text/markdown',
            'text/xml'
        ]
        return content_type in text_types or filename.lower().endswith(('.txt', '.html', '.htm', '.md', '.xml'))
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai texto de um arquivo de texto.
        
        Args:
            file_path: Caminho para o arquivo de texto
            
        Returns:
            Dict[str, Any]: Dicionário com texto extraído e metadados
        """
        try:
            # Determina o tipo de conteúdo
            content_type = mimetypes.guess_type(file_path)[0] or 'text/plain'
            
            # Tenta diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Não foi possível decodificar o arquivo de texto com os encodings disponíveis")
            
            # Se for HTML, extrai o texto
            if content_type == 'text/html':
                soup = BeautifulSoup(text, 'html.parser')
                text = soup.get_text(separator='\n')
            
            # Procura por padrões específicos no texto
            contract_numbers = self._extract_contract_numbers(text)
            school_names = self._extract_school_names(text)
            dates = self._extract_dates(text)
            
            return {
                'content_type': 'text/plain',
                'text': text,
                'metadata': {
                    'original_content_type': content_type,
                    'size': os.path.getsize(file_path)
                },
                'contract_numbers': contract_numbers,
                'school_names': school_names,
                'dates': dates
            }
        
        except Exception as e:
            logger.error(f"Erro ao processar arquivo de texto {file_path}: {str(e)}")
            return {
                'content_type': 'text/plain',
                'text': f"[Erro ao processar arquivo de texto: {str(e)}]",
                'metadata': {},
                'error': str(e)
            }
    
    # Reutiliza os métodos de extração do PDFProcessor
    def _extract_contract_numbers(self, text: str) -> List[str]:
        return PDFProcessor()._extract_contract_numbers(text)
    
    def _extract_school_names(self, text: str) -> List[str]:
        return PDFProcessor()._extract_school_names(text)
    
    def _extract_dates(self, text: str) -> List[str]:
        return PDFProcessor()._extract_dates(text)

class AttachmentProcessor:
    """Processador principal para anexos."""
    
    def __init__(self):
        """Inicializa o processador com os processadores disponíveis."""
        self.processors = [
            PDFProcessor(),
            DocxProcessor(),
            ExcelProcessor(),
            CSVProcessor(),
            TextProcessor()
        ]
    
    def process_attachment(self, attachment: Dict[str, Any]) -> Dict[str, Any]:
        """
        Processa um anexo.
        
        Args:
            attachment: Dicionário com informações do anexo
            
        Returns:
            Dict[str, Any]: Dicionário com conteúdo extraído e metadados
        """
        try:
            file_path = attachment.get('path', '')
            filename = attachment.get('filename', '')
            content_type = attachment.get('content_type', '')
            
            if not file_path or not os.path.exists(file_path):
                logger.error(f"Arquivo não encontrado: {file_path}")
                return {
                    'content_type': 'text/plain',
                    'text': "[Arquivo não encontrado]",
                    'metadata': {},
                    'error': "Arquivo não encontrado"
                }
            
            # Encontra o processador adequado
            for processor in self.processors:
                if processor.can_process(content_type, filename):
                    result = processor.process(file_path)
                    
                    # Adiciona informações do anexo ao resultado
                    result['filename'] = filename
                    result['original_content_type'] = content_type
                    
                    return result
            
            # Se nenhum processador for encontrado
            logger.warning(f"Nenhum processador encontrado para {content_type} ({filename})")
            return {
                'content_type': 'text/plain',
                'text': f"[Tipo de arquivo não suportado: {content_type}]",
                'metadata': {
                    'filename': filename,
                    'content_type': content_type
                },
                'error': "Tipo de arquivo não suportado"
            }
        
        except Exception as e:
            logger.error(f"Erro ao processar anexo: {str(e)}")
            return {
                'content_type': 'text/plain',
                'text': f"[Erro ao processar anexo: {str(e)}]",
                'metadata': {},
                'error': str(e)
            }

# Exemplo de uso
if __name__ == "__main__":
    processor = AttachmentProcessor()
    
    # Exemplo de anexo
    attachment = {
        'filename': 'exemplo.pdf',
        'content_type': 'application/pdf',
        'path': '/caminho/para/exemplo.pdf'
    }
    
    result = processor.process_attachment(attachment)
    print(f"Texto extraído: {result['text'][:100]}...")
    print(f"Números de contrato encontrados: {result.get('contract_numbers', [])}")
    print(f"Nomes de escolas encontrados: {result.get('school_names', [])}")
    print(f"Datas encontradas: {result.get('dates', [])}")
